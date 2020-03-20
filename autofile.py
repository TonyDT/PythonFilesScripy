# -*- coding: utf-8 -*-
  
import os,re,sys,shutil
import docx
from docx import Document
from docx.shared import Inches
#查询文件位置
def query():
    from docx import Document
    file_dir="/Users/taodong/Desktop/test/kuancheng/"
#    file_dir = r'C:\\Users\\jipaifafu\\Desktop\\jp\\hetong\\'
    L=[]
    for root, dirs, files in os.walk(file_dir):
        for file in files:
            
            if os.path.splitext(file)[1] == '.docx':
#                L.append(os.path.join(root, file))
#                print(os.path.join(root,file))#路径
#                print(file)#文件名
#                print(re.sub('.docx', '', file))#准备写入的名称
                name = re.sub('.docx', '', file)
#                print(file)
#                print(name)
                aa = (os.path.join(root, file))
#                print( aa)
                
                doc = Document(os.path.join(root, file))
                   #每一段的内容
                for para in doc.paragraphs:
                    print(para)
                    pass
#                        print(para.text)
                   
                   #每一段的编号、内容
                for i in range(len(doc.paragraphs)):
                        print(str(i), doc.paragraphs[i].text)
                        
                        if i == 1:
                             
                               ff = doc.paragraphs[1].text  + name
                               
                               print(len(ff))
                               doc.paragraphs[1].text = ff
                               doc.save(aa)
                               doc.paragraphs[1].text = "-"
                 
                L.append(file)
                
    return L
    
def reloadName(cityDate):
    print(cityDate)
#    srcDir = r'C:\\Users\\jipaifafu\\Desktop\\jp\\'
    srcDir = '/Users/taodong/Desktop/test/'
    dstDir = '/Users/taodong/Desktop/test/kuancheng/'
#    dstDir = r'C:\\Users\\jipaifafu\\Desktop\\jp\\hetong\\'
                
    #时间+序号
    files = os.listdir(dstDir)
    newName = cityDate + ".docx"
    
    for file in files:
        print('== %s' % file)
        #生成word文档
        srcFile=os.path.join(dstDir,file)
#        print('src ===>%s' % srcFile)
        dstFile=os.path.join(dstDir,newName)
#        print('pa ===>%s' % dstFile)
        os.rename(srcFile,dstFile)
         
        
if __name__=="__main__":

    while True:
                 cityDate = input('请输入你的城市编号和年月日：').strip()
                 reloadName(cityDate)
                 query()
             
    
    
    
